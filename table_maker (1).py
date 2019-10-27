import pandas as pd
import os
import docx
from docx import Document

builder = {'full Report' : '', 'investigators' : '', 'date' : '', 'complainant': '', 'report type' : '', 'reference': '', 'summary': '', 'names': []}

frame = pd.DataFrame.from_dict(builder)


# DO NOT TOUCH THIS!
def cellscraper(cell, header, rowlist, spot):
    if header in cell.paragraphs[0].text:
        try:
            rowlist[spot] = (cell.paragraphs[1].text.lower())
        except:
            rowlist[spot] = ''



# this is the scraper 
# DO NOT TOUCH!
def masterScraper(report, file):
    rowlist = ["","","","",""]
    rowlist[0] = (file.strip('.docx'))
    try:
        table = report.tables[0]
        for row in table.rows:
            for cell in row.cells:
                cellscraper(cell, 'Investigator', rowlist, 1)
                cellscraper(cell, 'Date', rowlist, 2)
                cellscraper(cell, 'Complainant', rowlist, 3)
                cellscraper(cell, 'Report Type', rowlist, 4)
    except:
        None
        
    nameList = []
    reference = ''
    summary = ''
    summaryExists = False
    referenceExists = False
    for section in report.paragraphs:
        if 'reference:' in section.text.lower():
            reference = (section.text.lower().lstrip('reference:').strip())
            referenceExists = True
            
        if referenceExists == False:
            if 'subject:' in section.text.lower():
                reference = (section.text.lower().lstrip('subject:').strip())
             
        if 'summary:' in section.text.lower():
            summary = (section.text.lower().lstrip('summary:').strip())
            summaryExists = True
            
        if summaryExists == False:
            if 'synopsis:' in section.text.lower():
                summary = (section.text.lower().lstrip('synopsis:').strip())
                
        if 'name:' in section.text.lower():
            name = (section.text.lstrip('Name:').strip().lower())
            if name == '':
                nameList.append('assosiated name(s)')
            else:
                nameList.append(name)
        
        
    rowlist.extend((reference, summary, ', '.join(nameList)))
    return(rowlist)

directory = os.fsdecode(r"X:\DCC Investigations")

paths = [x[0] for x in os.walk(directory)]
for path in paths:
    list_files = os.listdir(path)
    for file in list_files:
        if (file.endswith('.docx') or file.endswith('.doc')):
             try:
                report = docx.Document(path + r"\ ".strip() + file)
                frame.loc[len(frame)+1] = masterScraper(report, file)
             except:
                continue
frame = frame[frame.investigators != '']
frame = frame.reset_index().drop(labels = 'index', axis = 1)

frame.to_csv(r"X:\DCC Investigations\DCC Database\analysis\report data.csv")



import tkinter as tk 

root= tk.Tk() 
 
canvas1 = tk.Canvas(root, width = 300, height = 300)
canvas1.pack()

label1 = tk.Label(root, text='report data.csv and CLEMIS master.csv successfuly updated')
canvas1.create_window(150, 150, window=label1)

root.mainloop()
